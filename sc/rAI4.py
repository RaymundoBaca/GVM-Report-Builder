# rAI4.py
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import re
from PIL import Image
import natsort

def generate_report(txt_file, docx_file, img_dir, log_func=print, progress_callback=None):
    doc = Document()
    MAX_WIDTH = Inches(6.0)

    def set_run_style(run, bold=False, color=None, size=12):
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)

    def insert_image(path, scale=0.8, host=False):
        if not path.exists():
            log_func(f"⚠️ Image not found: {path}")
            return
        with Image.open(path) as img:
            w, h = img.size
            dpi = img.info.get("dpi", (96, 96))[0]
            if not dpi or dpi <= 0:
                dpi = 96
            width_in = (w / dpi) * scale
            height_in = (h / dpi) * scale
            if host:
                width_in *= 0.6
                height_in *= 0.6
            if width_in > MAX_WIDTH.inches:
                factor = MAX_WIDTH.inches / width_in
                width_in *= factor
                height_in *= factor
            doc.add_picture(str(path), width=Inches(width_in))
            doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def vuln_base_from_name(vuln_name: str) -> str:
        text = vuln_name.strip()
        text = re.sub(r'(High|Medium|Low)\s*\(CVSS:[^)]+\)', '', text, flags=re.IGNORECASE)
        text = re.sub(r'NVT:?', '', text, flags=re.IGNORECASE)
        text = text.replace("/", "").replace(".", "").replace(":", "")
        text = text.replace("'", "").replace('"', "")
        text = text.replace("(", "").replace(")", "")
        text = re.sub(r"[^\w\s\-]", " ", text)
        text = text.replace("_", " ")
        text = re.sub(r"\s+", " ", text).strip()
        words = text.split(" ")
        return "_".join(words) if words else "vuln"

    vuln_counters = {}
    def generate_unique_name(text: str) -> str:
        base = vuln_base_from_name(text)
        vuln_counters[base] = vuln_counters.get(base, 0) + 1
        return f"{base}{vuln_counters[base]}.png"

    def add_vuln_images_by_base(base_name: str):
        if not base_name:
            return
        idx, found_any = 1, False
        while True:
            path = img_dir / f"{base_name}{idx}.png"
            if not path.exists():
                break
            insert_image(path)
            found_any = True
            idx += 1
        if not found_any:
            p = doc.add_paragraph(f"[Image not found: {base_name}1.png]")
            set_run_style(p.runs[0], color=(0xAA, 0x00, 0x00))

    def add_repeated_vuln_image(vuln_name: str):
        if not vuln_name:
            return
        text = vuln_name.strip()
        text = text.replace("/", "").replace(".", "").replace(":", "")
        text = text.replace("'", "").replace('"', "")
        text = re.sub(r"[^\w\s\-]", " ", text)
        text = text.replace("_", " ")
        text = re.sub(r"\s+", " ", text).strip()
        words = text.split(" ")
        base_name = "_".join(words[:3]) if words else ""
        img_path = img_dir / f"{base_name}.png"
        if img_path.exists():
            insert_image(img_path)
        else:
            p = doc.add_paragraph(f"[Image not found: {base_name}.png]")
            set_run_style(p.runs[0], color=(0xAA, 0x00, 0x00))

    def add_cvss_line(paragraph_text):
        p = doc.add_paragraph()
        parts = re.split(r"(\(CVSS:\s*[\d\.]+\))", paragraph_text)
        cvss_m = re.search(r"\(CVSS:\s*([\d\.]+)\)", paragraph_text)
        shading = None
        if cvss_m:
            try:
                score = float(cvss_m.group(1))
                if 9.0 <= score <= 10.0:
                    shading = (0xEE, 0x00, 0x00)
                elif 7.0 <= score <= 8.9:
                    shading = (0xE9, 0x71, 0x32)
                elif 4.0 <= score <= 6.9:
                    shading = (0xFF, 0xC0, 0x00)
            except: pass
        for part in parts:
            run = p.add_run(part)
            if re.match(r"\(CVSS:\s*[\d\.]+\)", part) and shading:
                set_run_style(run, color=(0xFF,0xFF,0xFF))
                rPr = run._element.get_or_add_rPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), f'{shading[0]:02X}{shading[1]:02X}{shading[2]:02X}')
                shd.set(qn('w:val'), 'clear')
                rPr.append(shd)
            else:
                set_run_style(run)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    with open(txt_file, "r", encoding="utf-8") as f:
        lines = [ln.rstrip("\n") for ln in f]

    hosts_lines = [ln for ln in lines if ln.startswith("Host ")]
    total_hosts = len(hosts_lines)
    first_host = hosts_lines[0].split()[1] if hosts_lines else None
    network = ".".join(first_host.split(".")[:3]) + ".0/24" if first_host else "Unknown Network"

    # ===== Cover Page =====
    p = doc.add_paragraph()
    p.style = "Heading 4"
    run = p.add_run(network)
    set_run_style(run, bold=True, size=16)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = doc.add_paragraph("Below is a table of target hosts and the number of vulnerabilities found on each:")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(12)

    host_files = list(img_dir.glob("hosts*.png"))
    host_files_sorted = natsort.natsorted(host_files, key=lambda x: x.stem)
    for cover_img in host_files_sorted:
        insert_image(cover_img, host=True)

    p = doc.add_paragraph("To make this document more accessible and easy to understand, the following considerations will be applied:")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(12)

    for bullet in [
        "Each vulnerability will be described only the first time it appears; if it appears again, a note will mention it was already described.",
        "Low-risk vulnerabilities will be shown only when they have moderate to critical relevance."
    ]:
        para = doc.add_paragraph(bullet, style="List Bullet")
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para.runs[0].font.name = "Arial"
        para.runs[0].font.size = Pt(12)

    final_p = doc.add_paragraph(
        "The results of the analysis follow, explaining each host and the vulnerabilities found, affected ports, "
        "and mitigation suggestions highlighted in blue."
    )
    final_p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    final_p.runs[0].font.name = "Arial"
    final_p.runs[0].font.size = Pt(12)

    doc.add_page_break()

    # ===== Full line-by-line processing =====
    mitigation_mode = False
    current_host_idx = 0
    current_vuln_base = ""
    images_inserted_for_current_vuln = False
    host_counter = 0
    vulns_already_seen = set()  # Track vulnerabilities already described

    i = 0
    while i < len(lines):
        line = lines[i]

        # ===== Host line =====
        if line.startswith("Host "):
            host_counter += 1
            host = line.split()[1]
            current_host_idx += 1
            p = doc.add_paragraph(f"Host {host}")
            set_run_style(p.runs[0], bold=True, size=12)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            current_vuln_base = ""
            images_inserted_for_current_vuln = False
            mitigation_mode = False

            host_main = img_dir / f"host{current_host_idx}.png"
            if host_main.exists():
                insert_image(host_main, host=True)

            idx = 1
            while True:
                variant = img_dir / f"host{current_host_idx}.{idx}.png"
                if not variant.exists():
                    break
                insert_image(variant, host=True)
                idx += 1

            i += 1
            continue

        # ===== Port summary image =====
        if line.strip() == "The following image shows a summary of the affected ports for this particular host.":
            p = doc.add_paragraph(line)
            set_run_style(p.runs[0])
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            i += 1
            continue

        # ===== Vulnerability line =====
        if line.startswith("The following vulnerability is called NVT:"):
            vuln_full = line.replace("The following vulnerability is called NVT:", "").strip()
            add_cvss_line(f"The following vulnerability is called NVT: {vuln_full}")
            vuln_name_no_cvss = re.sub(r"\(CVSS:\s*[\d\.]+\)\s*$", "", vuln_full).strip()
            current_vuln_base = vuln_base_from_name(vuln_name_no_cvss)
            images_inserted_for_current_vuln = False
            i += 1
            continue

        # ===== Description =====
        if line.startswith("Description:"):
            if vuln_name_no_cvss not in vulns_already_seen:
                add_vuln_images_by_base(current_vuln_base)  # Insert numbered images _1.png, _2.png...
                vulns_already_seen.add(vuln_name_no_cvss)
                images_inserted_for_current_vuln = True
            p = doc.add_paragraph(line)
            set_run_style(p.runs[0])
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            i += 1
            continue

        # ===== Already described =====
        if line.startswith("This vulnerability has already been described previously"):
            if current_vuln_base:
                # 🔹 Tomar solo las 3 primeras palabras para la imagen sin numeración
                words = current_vuln_base.split("_")
                base_name_repeat = "_".join(words[:3])
                main_img_path = img_dir / f"{base_name_repeat}.png"
                if main_img_path.exists():
                    insert_image(main_img_path)
                else:
                    p = doc.add_paragraph(f"[Image not found: {base_name_repeat}.png]")
                    set_run_style(p.runs[0], color=(0xAA, 0x00, 0x00))

            p = doc.add_paragraph(line)
            set_run_style(p.runs[0])
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            i += 1
            continue

        # ===== Mitigation =====
        if line.startswith("Mitigation:"):
            mitigation_mode = True
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_run_style(run, color=(0x4C, 0x94, 0xD8))
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            i += 1
            continue

        if mitigation_mode:
            end_of_mitigation = (
                line.startswith("The following vulnerability is called NVT:") or
                line.startswith("Host ") or
                line.strip() == ""
            )
            if end_of_mitigation:
                mitigation_mode = False
                if host_counter < total_hosts:
                    doc.add_page_break()
                continue
            else:
                p = doc.add_paragraph()
                run = p.add_run(line)
                set_run_style(run, color=(0x4C, 0x94, 0xD8))
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                i += 1
                continue

        # ===== Blank line =====
        if line.strip() == "":
            i += 1
            continue

        # ===== Default text =====
        p = doc.add_paragraph(line)
        set_run_style(p.runs[0])
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        i += 1

    docx_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_file)
    log_func(f"✅ Report generated: {docx_file}")

def main(pdf_path, txt_path, log, output_dir=None, progress_callback=None):
    """
    pdf_path: ignored here, just for GUI compatibility
    txt_path: ignored, Summary.txt is used directly
    log: logging function
    output_dir: folder where the report will be generated
    progress_callback: optional callback(float) to report progress (0.0 to 1.0)
    """
    if output_dir:
        results_dir = Path(output_dir) / "Results"
    else:
        results_dir = Path(__file__).resolve().parent / "Results"

    summary_txt = results_dir / "Summary.txt"
    if not summary_txt.exists():
        log(f"⛔ File not found: {summary_txt}")
        return

    input_dir = results_dir / "tables_png"
    if not input_dir.exists():
        log(f"⛔ Folder not found: {input_dir}")
        return

    if progress_callback:
        progress_callback(0.0)

    results_dir.mkdir(parents=True, exist_ok=True)
    docx_file = results_dir / "VulnReport.docx"

    generate_report(summary_txt, docx_file, input_dir, log, progress_callback=progress_callback)

    if progress_callback:
        progress_callback(1.0)
